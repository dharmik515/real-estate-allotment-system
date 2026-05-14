from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

class AllotmentDocumentGenerator:
    def __init__(self):
        pass
    
    def format_currency(self, amount):
        """Format amount in Indian currency style"""
        amount_str = str(int(amount))
        if len(amount_str) <= 3:
            return amount_str
        
        last_three = amount_str[-3:]
        remaining = amount_str[:-3]
        
        formatted = ""
        while len(remaining) > 2:
            formatted = "," + remaining[-2:] + formatted
            remaining = remaining[:-2]
        
        if remaining:
            formatted = remaining + formatted
        
        return formatted + "," + last_three
    
    def number_to_words_indian(self, num):
        """Convert number to words in Indian format"""
        ones = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine"]
        tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
        teens = ["Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
        
        def convert_below_thousand(n):
            if n == 0:
                return ""
            elif n < 10:
                return ones[n]
            elif n < 20:
                return teens[n - 10]
            elif n < 100:
                return tens[n // 10] + (" " + ones[n % 10] if n % 10 != 0 else "")
            else:
                return ones[n // 100] + " Hundred" + (" " + convert_below_thousand(n % 100) if n % 100 != 0 else "")
        
        if num == 0:
            return "Zero"
        
        crore = num // 10000000
        num %= 10000000
        lakh = num // 100000
        num %= 100000
        thousand = num // 1000
        num %= 1000
        
        result = []
        
        if crore:
            result.append(convert_below_thousand(crore) + " Crore")
        if lakh:
            result.append(convert_below_thousand(lakh) + " Lakh")
        if thousand:
            result.append(convert_below_thousand(thousand) + " Thousand")
        if num:
            result.append(convert_below_thousand(num))
        
        return " ".join(result) + " Only"
    
    def generate_allotment_letter(self, project_data, allotment_data, allottees_data, output_path):
        """Generate allotment letter in DOCX format"""
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run("ALLOTMENT LETTER")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Allotment Number and Date
        header = doc.add_paragraph()
        header.add_run(f"No. {allotment_data['allotment_number']}").font.size = Pt(11)
        header.add_run(" " * 80)
        header.add_run(f"Date: {allotment_data['issue_date']}").font.size = Pt(11)
        
        doc.add_paragraph()
        
        # To section - Allottees
        doc.add_paragraph("To,")
        for allottee in allottees_data:
            name_para = doc.add_paragraph(allottee['name'])
            name_para.runs[0].bold = True
        
        # Address
        address_lines = allottees_data[0]['address'].split('\n')
        for line in address_lines:
            doc.add_paragraph(line)
        
        doc.add_paragraph(f"Phone no.: {allottees_data[0]['phone']}")
        
        # PAN and Aadhar
        pan_cards = " / ".join([a['pan_card'] for a in allottees_data if a['pan_card']])
        aadhar_cards = " / ".join([a['aadhar_card'] for a in allottees_data if a['aadhar_card']])
        
        doc.add_paragraph(f"Pan Card No: {pan_cards}")
        doc.add_paragraph(f"Aadhar Card No: {aadhar_cards}")
        doc.add_paragraph(f"Email id: {allottees_data[0]['email']}")
        
        doc.add_paragraph()
        
        # Subject
        subject = doc.add_paragraph()
        subject.add_run(f"Sub: - Your request for allotment of commercial premises in the project known as {project_data['project_name']} having MahaRERA Registration No. {project_data['mahrera_no']}.").bold = True
        
        doc.add_paragraph()
        doc.add_paragraph("Sir/Madam,")
        
        # Section 1: Allotment of the said unit
        heading1 = doc.add_paragraph()
        heading1.add_run("Allotment of the said unit:").bold = True
        
        allotment_text = f"This has reference to your request referred at the above subject. In that regard, we have the pleasure to inform that you have been allotted a {allotment_data['unit_type'].lower()} premises bearing {allotment_data['unit_type']} No. {allotment_data['unit_number']} admeasuring RERA carpet area {allotment_data['carpet_area_sqm']:.2f} sq.mtrs equivalent to {allotment_data['carpet_area_sqft']:.0f} sq.ft. situated on {allotment_data['floor']} floor in {allotment_data['wing']} wing in the project known as {project_data['project_name']}, having MahaRERA Registration No. {project_data['mahrera_no']}, hereinafter referred to as \"the said unit\", being developed on land bearing {project_data['survey_no']}, {project_data['hissa_no']} and {project_data['cts_no']} of {project_data['location']} and now {project_data['final_plot_no']}, {project_data['tps_details']} admeasuring about {project_data['plot_area']} lying and situate at {project_data['location']} for a total consideration of Rs. {self.format_currency(allotment_data['total_consideration'])}/- (Rupees {self.number_to_words_indian(int(allotment_data['total_consideration']))}) exclusive of GST, Stamp Duty and Registration Charges."
        
        doc.add_paragraph(allotment_text)
        
        # Section 2: Parking
        heading2 = doc.add_paragraph()
        heading2.add_run("Allotment of parking space(s):").bold = True
        
        if allotment_data['parking_required'] == 'Yes':
            parking_text = f"Allottee has been allotted {allotment_data['parking_type']} parking at {allotment_data['parking_location']} ({allotment_data['parking_size']} size) in the said project."
        else:
            parking_text = "Allottee has informed the promoter that he/she does not require any car parking space in said project. Accordingly, no reservation of car parking is made against said Unit. Allottee undertakes, assures and guarantees not to claim any car parking space in said project in future, nor raise any objection to use of car parking by other Allottees"
        
        doc.add_paragraph(parking_text)
        
        # Section 3: Receipt of part consideration
        heading3 = doc.add_paragraph()
        heading3.add_run("Receipt of part consideration:").bold = True
        
        receipt_text = f"We confirm to have received from you an amount of Rs. {self.format_currency(allotment_data['booking_amount'])}/- (Rupees {self.number_to_words_indian(int(allotment_data['booking_amount']))}), of the total consideration value of the said unit as booking amount/advance payment on various dates, through cheques/RTGS/NEFT."
        
        doc.add_paragraph(receipt_text)
        
        # Section 4: Disclosures
        heading4 = doc.add_paragraph()
        heading4.add_run("Disclosures of information:").bold = True
        
        doc.add_paragraph("I/we have made available to you the following information namely: -")
        doc.add_paragraph("The sanctioned plans, layout plans, along with specifications, approved by the competent authority are displayed at the project site and has also been uploaded on MahaRERA website.")
        doc.add_paragraph("The website address of MahaRERA is https://mahrera.mahaonline.gov.in/#")
        
        doc.add_paragraph()
        
        # Section 5: Encumbrances
        heading5 = doc.add_paragraph()
        heading5.add_run("Encumbrances:").bold = True
        
        doc.add_paragraph("I/ we hereby confirm that the said unit is free from all encumbrances and I/ we hereby further confirm that no encumbrances shall be created on the said unit.")
        
        # Section 6: Further payments
        heading6 = doc.add_paragraph()
        heading6.add_run("6.      Further payments:").bold = True
        
        doc.add_paragraph("Further payments towards consideration of the said unit as well as of the car parking space(s) shall be made by you, in manner and at times as well as on the terms and conditions as more specifically enumerated / stated in the agreement for sale to be entered into between ourselves and yourselves.")
        
        # Section 7: Possession
        heading7 = doc.add_paragraph()
        heading7.add_run("Possession:").bold = True
        
        possession_text = f"The said unit along with the garage(s)/ covered car parking spaces (s) shall be handed over to you on or before {allotment_data['possession_date']} subject to the payment of the consideration amount of the said unit as well as of the covered car parking space(s) in the manner and at times as well as per the terms and conditions as more specifically enumerated/ stated in the agreement for sale to be entered between ourselves and yourselves."
        
        doc.add_paragraph(possession_text)
        
        # Section 8: Interest Payment
        heading8 = doc.add_paragraph()
        heading8.add_run("Interest Payment:").bold = True
        
        doc.add_paragraph("In case of delay in making any payments, you shall be liable to pay interest at the rate which shall be the State Bank of India highest Marginal cost of Lending Rate plus two percent.")
        
        # Section 9: Cancellation
        heading9 = doc.add_paragraph()
        heading9.add_run("Cancellation of Allotment:").bold = True
        
        doc.add_paragraph("In case you desire to cancel the booking an amount mentioned in the table hereunder written* would be deducted and the balance amount due and payable shall be refunded to you without interest within 45 days from the date of receipt of your letter requesting to cancel the said booking.")
        
        doc.add_paragraph("*The amount deducted shall not exceed the amount as mentioned in the table above.")
        
        doc.add_paragraph("In the event the amount due and payable referred in Clause 9 i) above is not refunded within 45 days from the date of receipt of your letter requesting to cancel the said booking, you shall be entitled to receive the balance amount with interest calculated at the rate which shall be the State Bank of India highest Marginal Cost of lending Rate plus two percent.")
        
        # Continue with remaining sections...
        # Section 10: Other payments
        heading10 = doc.add_paragraph()
        heading10.add_run("Other payments:").bold = True
        
        gst_text = f"You shall make the payment of GST (Rs. {self.format_currency(allotment_data['gst_amount'])}/-), stamp Duty and registration charges, as applicable and such other payments as more specifically mentioned in the agreement for sale, the proforma whereof is enclosed herewith in terms of Clause 11 hereunder written."
        
        doc.add_paragraph(gst_text)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Remaining standard clauses (11-15)
        self._add_standard_clauses(doc)
        
        # Cancellation Table
        doc.add_paragraph()
        self._add_cancellation_table(doc, project_data, allotment_data)
        
        # Bank Details Schedule
        doc.add_paragraph()
        self._add_bank_details(doc, project_data, allottees_data[0])
        
        # Signature section
        doc.add_paragraph()
        doc.add_paragraph()
        
        sig_para = doc.add_paragraph()
        sig_para.add_run(f"For {project_data['promoter_company']}").bold = True
        
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        auth_sig = doc.add_paragraph("Authorized Signatory")
        auth_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Confirmation section
        doc.add_paragraph()
        confirm_heading = doc.add_paragraph()
        confirm_heading.add_run("CONFIRMATION AND ACKNOWLEDGEMENT").bold = True
        confirm_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("I/we have read and understood the contents of this allotment letter and the Annexure. I/we hereby agree and accept the terms and conditions as stipulated in this allotment letter.")
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Allottee signatures
        for allottee in allottees_data:
            doc.add_paragraph()
            doc.add_paragraph()
            name_sig = doc.add_paragraph(allottee['name'])
            name_sig.runs[0].bold = True
        
        doc.add_paragraph()
        doc.add_paragraph(" (Allottee/s)")
        doc.add_paragraph(f"Date: {allotment_data['issue_date']}")
        doc.add_paragraph(f"Place:  Mumbai")
        
        # Save document
        doc.save(output_path)
        return output_path
    
    def _add_standard_clauses(self, doc):
        """Add standard clauses 11-15"""
        
        # Clause 11
        heading11 = doc.add_paragraph()
        heading11.add_run("Proforma of the agreement for sale and binding effect:").bold = True
        doc.add_paragraph("The preforma of the agreement for sale to be entered into between ourselves and yourselves is enclosed herewith for your ready reference. Forwarding the proforma of the agreement for sale does not create a binding obligation on the part of ourselves and yourselves until compliance by yourselves of the mandate as stated in Clause 12.")
        
        # Clause 12
        heading12 = doc.add_paragraph()
        heading12.add_run("Execution and registration of the agreement for sale;").bold = True
        doc.add_paragraph("You shall execute the agreement for sale and appear for registration of the same before the concerned Sub-Registrar within a period of 2 months from the date of issuance of this letter or within such period as may be communicated to you. The said period of 2 months can be further extended on our mutual understanding")
        
        doc.add_paragraph("In the event the booking amount is collected in stages and if the allottee fails to pay the subsequent stage installment, the promoter shall serve upon the allottee a notice calling upon the allottee to pay the subsequent stage installment within 15 (fifteen) days which if not complied, the promoter shall be entitled to cancel this allotment letter.")
        
        doc.add_paragraph("On cancellation of the allotment letter the promoter shall be entitled to forfeit the amount paid by the allottee or such amount as mentioned in the Table enumerated in Clause 9 whichever is less. In no event the amount to be forfeited shall exceed the amount mentioned in the above referred Table. Except for the above all the terms and conditions as enumerated in this allotment letter shall be applicable even for cases where booking amount is collected in stages. In event of cancellation, amount paid by you, after deductions of forfeiture charges, shall be refunded to your Bank Account as mentioned in the Schedule I hereunder written and thereafter, you shall not have any right, title, claim and interest over the Said Flat and I/we shall be entitled to dispose of the same as I/we deem fit and proper including selling the said flat to any third party.")
        
        doc.add_paragraph("If you fail to execute the agreement for sale and appear for registration of the same before the concerned Sub-Registrar within the stipulated period 2 months from the date of issuance of this letter or within such period as may be communicated to you, I/we shall be entitled to serve upon you a notice calling upon you to execute the agreement for sale and appear for registration of the same within 15 (fifteen) days, which if not complied, I/we shall be entitled to cancel this allotment letter and further I/we shall be entitled to forfeit an amount not exceeding 2% of the cost of the said unit and the balance amount if any due and payable shall be refunded without interest within 45 days from the date of expiry of the notice period.")
        
        doc.add_paragraph("In the event the balance amount due and payable refereed in Clause 12 ii) above is not refunded within 45 days from the date of expiry of the notice period, you shall be entitled to receive the balance amount with interest calculated at the rate which shall be the State Bank of India highest Marginal cost of Lending Rate plus two percent")
        
        # Clause 13
        heading13 = doc.add_paragraph()
        heading13.add_run("Validity of allotment letter:").bold = True
        doc.add_paragraph("This allotment letter shall not be construed to limit your rights and interest upon execution and registration of the agreement for sale between ourselves and yourselves. Cancellation of allotment of the said unit thereafter, shall be covered by the terms and conditions of the said registered document.")
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Clause 14
        heading14 = doc.add_paragraph()
        heading14.add_run("Headings:").bold = True
        doc.add_paragraph("Headings are inserted for convenience only and shall not affect the construction of the various Clauses of this allotment letter.")
    
    def _add_cancellation_table(self, doc, project_data, allotment_data):
        """Add cancellation charges table"""
        import json
        
        cancellation_data = json.loads(project_data['cancellation_table'])
        
        if not cancellation_data:
            # Default table
            cancellation_data = [
                {"period": "within 15 days from issuance of the allotment letter;", "deduction": "Nil"},
                {"period": "within 16 to 30 days from issuance of the allotment letter;", "deduction": "1% of the cost of the said unit;"},
                {"period": "within 31 to 60 days from issuance of the allotment letter;", "deduction": "1.5 % of the cost of the said unit;"},
                {"period": "after 61 days from issuance of the allotment letter", "deduction": "2% of the cost of the said unit."}
            ]
        
        table = doc.add_table(rows=len(cancellation_data) + 1, cols=3)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Sr. No."
        header_cells[1].text = "If the letter requesting to cancel the booking is received"
        header_cells[2].text = "Amount to be deducted"
        
        # Data rows
        for idx, row_data in enumerate(cancellation_data, 1):
            row_cells = table.rows[idx].cells
            row_cells[0].text = str(idx)
            row_cells[1].text = row_data['period']
            row_cells[2].text = row_data['deduction']
    
    def _add_bank_details(self, doc, project_data, allottee_data):
        """Add bank details schedule"""
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = "SCHEDULE I"
        header_cells[1].text = "SCHEDULE I"
        
        # Subheader
        subheader_cells = table.rows[1].cells
        subheader_cells[0].text = "Promoter's Bank Details"
        subheader_cells[1].text = "Allottee's Bank Details"
        
        # Account Name
        row2_cells = table.rows[2].cells
        row2_cells[0].text = f"Account Name: {project_data['promoter_account_name']}"
        row2_cells[1].text = f"Account Name: {allottee_data['account_name']}"
        
        # Account No
        row3_cells = table.rows[3].cells
        row3_cells[0].text = f"Account No.: {project_data['promoter_account_no']}"
        row3_cells[1].text = f"Account No. {allottee_data['account_no']}"
        
        # Bank Name
        row4_cells = table.rows[4].cells
        row4_cells[0].text = f"Bank Name: {project_data['promoter_bank_name']}"
        row4_cells[1].text = f"Bank Name: {allottee_data['bank_name']}"
        
        # Branch and IFSC
        row5_cells = table.rows[5].cells
        row5_cells[0].text = f"Branch Name: {project_data['promoter_branch']}\nIFSC Code: {project_data['promoter_ifsc']}"
        row5_cells[1].text = f"Branch Name: {allottee_data['branch']}\nIFSC Code:  {allottee_data['ifsc']}"
