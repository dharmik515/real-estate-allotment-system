from docx import Document
import sys

# Read the document
doc = Document('ALLOTMENT NO. 6 Shop no. 4 Kavya harish suvarna - Copy.docx')

# Extract all paragraphs
for para in doc.paragraphs:
    print(para.text)

# Extract tables if any
print("\n\n=== TABLES ===\n")
for table in doc.tables:
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        print(" | ".join(row_data))
    print("\n")
